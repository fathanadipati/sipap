#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Final Report untuk AureliaBox - SIPAP
Script untuk membuat dokumen Word (.docx) berisi laporan komprehensif
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_title(doc, text):
    """Tambah judul"""
    title = doc.add_paragraph(text)
    title.style = 'Heading 1'
    title_format = title.runs[0]
    title_format.font.size = Pt(24)
    title_format.font.bold = True
    title_format.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return title

def add_heading(doc, text, level=2):
    """Tambah heading"""
    heading = doc.add_paragraph(text)
    heading.style = f'Heading {level}'
    heading_format = heading.runs[0]
    heading_format.font.size = Pt(14 if level == 2 else 12)
    heading_format.font.bold = True
    heading_format.font.color.rgb = RGBColor(0, 102, 204)
    return heading

def add_normal_text(doc, text):
    """Tambah teks normal"""
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet_point(doc, text):
    """Tambah bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = 1.3
    return p

def add_table_to_doc(doc, headers, rows):
    """Tambah tabel ke dokumen"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.runs[0].font.bold = True
    
    # Rows
    for i, row in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, cell_text in enumerate(row):
            row_cells[j].text = str(cell_text)
    
    return table

def create_report():
    """Buat laporan akhir"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ========== COVER PAGE ==========
    add_title(doc, "LAPORAN AKHIR APLIKASI")
    doc.add_paragraph()
    
    # Logo/Nama Aplikasi
    title = doc.add_paragraph("AureliaBox")
    title_run = title.runs[0]
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Smart Package Management System")
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n\n")
    
    # Info Dokumen
    info_table = doc.add_table(rows=4, cols=2)
    info_table.autofit = False
    
    info_data = [
        ("Nama Proyek", "AureliaBox - Smart Package Management System"),
        ("Tanggal Laporan", datetime.now().strftime("%d %B %Y")),
        ("Status", "Production Ready"),
        ("Versi", "1.0")
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        for cell in info_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.3
    
    doc.add_page_break()
    
    # ========== DAFTAR ISI ==========
    add_title(doc, "DAFTAR ISI")
    
    toc_items = [
        "1. Pendahuluan",
        "2. Latar Belakang & Tujuan",
        "3. Fitur Utama",
        "4. Teknologi yang Digunakan",
        "5. Rancangan Sistem",
        "6. Struktur Database",
        "7. Struktur Folder & File",
        "8. Panduan Instalasi",
        "9. Panduan Penggunaan",
        "10. Keamanan Sistem",
        "11. Kesimpulan"
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # ========== 1. PENDAHULUAN ==========
    add_heading(doc, "1. PENDAHULUAN", 1)
    add_normal_text(doc, 
        "AureliaBox adalah sebuah sistem manajemen paket yang dirancang khusus untuk "
        "mengelola alur pengiriman paket di apartemen premium. Sistem ini dikembangkan "
        "dengan tujuan memberikan solusi terpadu yang efisien dan user-friendly dalam "
        "menangani penerimaan, penyimpanan, dan pengambilan paket oleh penghuni.")
    
    add_normal_text(doc,
        "Dikembangkan khusus untuk THE GRAND AURELIA RESIDENCE, aplikasi ini memadukan "
        "teknologi modern dengan desain antarmuka yang intuitif untuk meningkatkan "
        "pengalaman pengguna dan efisiensi operasional.")
    
    doc.add_page_break()
    
    # ========== 2. LATAR BELAKANG & TUJUAN ==========
    add_heading(doc, "2. LATAR BELAKANG & TUJUAN", 1)
    
    add_heading(doc, "2.1 Latar Belakang", 2)
    add_normal_text(doc,
        "Setiap hari, apartemen premium menerima ratusan paket dari berbagai ekspedisi. "
        "Sistem penerimaan paket yang manual atau tidak terstruktur sering kali menyebabkan:")
    
    problems = [
        "Kesalahan pencatatan data paket",
        "Penghuni tidak mengetahui jika paketnya sudah tiba",
        "Kesulitan dalam melacak status paket",
        "Inefisiensi waktu penerimaan dan pengambilan paket",
        "Kurangnya dokumentasi yang baik untuk keperluan audit"
    ]
    
    for problem in problems:
        add_bullet_point(doc, problem)
    
    add_heading(doc, "2.2 Tujuan Pengembangan", 2)
    
    goals = [
        "Menciptakan sistem manajemen paket yang terintegrasi dan efisien",
        "Memberikan notifikasi real-time kepada penghuni saat paket tiba",
        "Menyediakan dashboard monitoring untuk admin dan resepsionis",
        "Mengotomatisasi proses penerimaan, penyimpanan, dan pengambilan paket",
        "Menjaga keamanan data dan mengimplementasikan kontrol akses berbasis role",
        "Memberikan user experience yang intuitif dan responsif"
    ]
    
    for goal in goals:
        add_bullet_point(doc, goal)
    
    doc.add_page_break()
    
    # ========== 3. FITUR UTAMA ==========
    add_heading(doc, "3. FITUR UTAMA", 1)
    
    features = {
        "3.1 Sistem Autentikasi & Role Management": [
            "Login dengan username dan password yang aman",
            "Password hashing menggunakan bcrypt untuk keamanan maksimal",
            "Session management yang robust dan aman",
            "RBAC (Role-Based Access Control) dengan 3 role: Admin, Resepsionis, Penghuni",
            "Proteksi akses langsung untuk halaman yang tidak tersedia"
        ],
        "3.2 Dashboard": [
            "Dashboard khusus untuk setiap role (Admin, Resepsionis, Penghuni)",
            "Statistik paket real-time dengan visualisasi data",
            "Quick access button ke fitur utama",
            "Informasi ringkas yang relevan untuk setiap pengguna",
            "Responsive design untuk semua ukuran layar"
        ],
        "3.3 Manajemen Data Penghuni (Admin)": [
            "CRUD lengkap untuk data penghuni",
            "Penyimpanan informasi unit, kontak, dan kontak darurat",
            "Integrasi otomatis dengan sistem akun pengguna",
            "Validasi data yang ketat",
            "Cascade delete untuk menjaga integritas data"
        ],
        "3.4 Manajemen Paket (Admin & Resepsionis)": [
            "Penerimaan paket baru dengan form lengkap",
            "Pencatatan detail: kurir, ekspedisi, jenis paket, dan deskripsi",
            "Auto-generate nomor paket unik dan otomatis",
            "Penyimpanan paket ke loker dengan nomor unik",
            "Update status paket: Diterima → Disimpan → Diambil",
            "Filter dan pencarian paket yang powerful",
            "Edit dan delete paket dengan konfirmasi"
        ],
        "3.5 Sistem Notifikasi Real-Time": [
            "Notifikasi otomatis terkirim ke penghuni saat paket tiba",
            "Bell icon di navbar dengan badge counter",
            "Daftar notifikasi dengan status baca/belum baca",
            "Refresh real-time setiap 5 detik menggunakan AJAX",
            "Mark as read/unread untuk setiap notifikasi",
            "API endpoint untuk integrasi dengan sistem lain"
        ],
        "3.6 Profil Pengguna": [
            "Edit profil (nama, email)",
            "Lihat informasi akun yang lengkap",
            "Manajemen password",
            "Update real-time ke database"
        ]
    }
    
    for section, feature_list in features.items():
        add_heading(doc, section, 2)
        for feature in feature_list:
            add_bullet_point(doc, feature)
    
    doc.add_page_break()
    
    # ========== 4. TEKNOLOGI YANG DIGUNAKAN ==========
    add_heading(doc, "4. TEKNOLOGI YANG DIGUNAKAN", 1)
    
    tech_table_data = [
        ("Backend", "PHP 7.4+ (Native - Tanpa Framework)"),
        ("Database", "MySQL 5.7+"),
        ("Frontend", "HTML5, CSS3, Vanilla JavaScript"),
        ("CSS Framework", "Bootstrap 5"),
        ("Icon", "Bootstrap Icons"),
        ("Server", "Apache (XAMPP)"),
        ("AJAX", "Vanilla JavaScript + XMLHttpRequest"),
        ("Password Hashing", "bcrypt"),
        ("Version Control", "Git")
    ]
    
    add_table_to_doc(doc, ["Komponen", "Teknologi"], tech_table_data)
    
    add_heading(doc, "4.1 Alasan Pemilihan Teknologi", 2)
    
    reasons = [
        ("PHP Native", "Lightweight, mudah di-deploy, tidak perlu dependency kompleks, cocok untuk XAMPP"),
        ("MySQL", "Reliable, widely used, support untuk relational database dengan foreign key"),
        ("Bootstrap 5", "Modern UI components, responsive grid system, extensive documentation"),
        ("Vanilla JavaScript", "No dependencies, fast performance, native browser support untuk AJAX"),
        ("Apache", "Standard web server, integrated dengan XAMPP, easy configuration")
    ]
    
    for tech, reason in reasons:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f"{tech}: ")
        run.bold = True
        p.add_run(reason)
    
    doc.add_page_break()
    
    # ========== 5. RANCANGAN SISTEM ==========
    add_heading(doc, "5. RANCANGAN SISTEM", 1)
    
    add_heading(doc, "5.1 Arsitektur Aplikasi", 2)
    add_normal_text(doc,
        "AureliaBox menggunakan arsitektur 3-Tier (Presentation-Business Logic-Data) yang "
        "klasik dan proven:")
    
    doc.add_paragraph("Presentation Layer (Frontend)", style='List Bullet')
    doc.add_paragraph("HTML, CSS, JavaScript untuk user interface", style='List Bullet 2')
    
    doc.add_paragraph("Business Logic Layer (Backend)", style='List Bullet')
    doc.add_paragraph("PHP files yang menghandle business logic", style='List Bullet 2')
    
    doc.add_paragraph("Data Access Layer (Database)", style='List Bullet')
    doc.add_paragraph("MySQL untuk persistent data storage", style='List Bullet 2')
    
    add_heading(doc, "5.2 Workflow Paket Masuk", 2)
    add_normal_text(doc, "Berikut adalah workflow lengkap saat paket masuk ke sistem:")
    
    workflow_steps = [
        "Kurir datang dengan paket",
        "Resepsionis klik menu 'Terima Paket Baru'",
        "Resepsionis mengisi form: nama pengirim, ekspedisi, jenis paket, deskripsi",
        "Resepsionis memilih unit/penghuni penerima paket",
        "Resepsionis memasukkan nomor loker tempat paket disimpan",
        "Resepsionis klik tombol 'Simpan'",
        "Sistem auto-generate nomor paket unik",
        "Database update dengan status 'Disimpan'",
        "Sistem otomatis mengirim notifikasi ke penghuni",
        "Penghuni menerima notifikasi real-time di dashboard",
        "Penghuni dapat melihat detail paket dan lokernya"
    ]
    
    for i, step in enumerate(workflow_steps, 1):
        doc.add_paragraph(f"{i}. {step}", style='List Number')
    
    add_heading(doc, "5.3 User Roles & Permissions", 2)
    
    roles_data = [
        ("Admin", "Manage users, manage residents, manage packages, view all data, access admin panel"),
        ("Resepsionis", "Receive packages, manage packages (edit/delete), view resident data, view notifications"),
        ("Penghuni", "View own packages, view notifications, edit profile, view own data")
    ]
    
    add_table_to_doc(doc, ["Role", "Permissions"], roles_data)
    
    doc.add_page_break()
    
    # ========== 6. STRUKTUR DATABASE ==========
    add_heading(doc, "6. STRUKTUR DATABASE", 1)
    
    add_heading(doc, "6.1 Entity Relationship Diagram (Konsep)", 2)
    add_normal_text(doc,
        "Database terdiri dari 4 tabel utama yang saling berelasi untuk menyimpan data pengguna, "
        "penghuni, paket, dan notifikasi:")
    
    # Tabel Users
    add_heading(doc, "6.2 Tabel: users", 2)
    users_data = [
        ("id", "INT", "Primary Key, Auto Increment"),
        ("username", "VARCHAR(50)", "UNIQUE, untuk login"),
        ("email", "VARCHAR(100)", "UNIQUE, email pengguna"),
        ("password", "VARCHAR(255)", "Hashed dengan bcrypt"),
        ("role", "ENUM", "admin, resepsionis, atau penghuni"),
        ("nama_lengkap", "VARCHAR(100)", "Nama lengkap pengguna"),
        ("is_active", "BOOLEAN", "Status aktif/nonaktif"),
        ("created_at", "TIMESTAMP", "Waktu pembuatan record"),
        ("updated_at", "TIMESTAMP", "Waktu update terakhir")
    ]
    add_table_to_doc(doc, ["Field", "Tipe", "Keterangan"], users_data)
    
    # Tabel Penghuni
    add_heading(doc, "6.3 Tabel: penghuni", 2)
    penghuni_data = [
        ("id", "INT", "Primary Key, Auto Increment"),
        ("user_id", "INT", "Foreign Key ke tabel users"),
        ("nomor_unit", "VARCHAR(20)", "UNIQUE, nomor unit apartemen"),
        ("blok", "VARCHAR(10)", "Blok unit"),
        ("lantai", "INT", "Lantai unit"),
        ("nomor_hp", "VARCHAR(15)", "Nomor telepon penghuni"),
        ("nama_kontak_darurat", "VARCHAR(100)", "Nama kontak darurat"),
        ("nomor_kontak_darurat", "VARCHAR(15)", "Nomor kontak darurat"),
        ("created_at", "TIMESTAMP", "Waktu pembuatan record"),
        ("updated_at", "TIMESTAMP", "Waktu update terakhir")
    ]
    add_table_to_doc(doc, ["Field", "Tipe", "Keterangan"], penghuni_data)
    
    # Tabel Paket
    add_heading(doc, "6.4 Tabel: paket", 2)
    paket_data = [
        ("id", "INT", "Primary Key, Auto Increment"),
        ("nomor_paket", "VARCHAR(50)", "UNIQUE, auto-generated identifier"),
        ("penghuni_id", "INT", "Foreign Key ke tabel penghuni"),
        ("nama_pengirim", "VARCHAR(100)", "Nama pengirim paket"),
        ("nama_kurir", "VARCHAR(100)", "Nama kurir yang mengantarkan"),
        ("nama_ekspedisi", "VARCHAR(100)", "Nama perusahaan ekspedisi"),
        ("jenis_paket", "VARCHAR(50)", "Kategori paket"),
        ("deskripsi", "TEXT", "Deskripsi isi paket"),
        ("nomor_loker", "VARCHAR(20)", "Nomor loker penyimpanan"),
        ("status", "ENUM", "diterima, disimpan, atau diambil"),
        ("resepsionis_id", "INT", "Foreign Key ke user (resepsionis)"),
        ("tanggal_terima", "DATETIME", "Waktu paket diterima"),
        ("tanggal_diambil", "DATETIME", "Waktu paket diambil (nullable)"),
        ("catatan", "TEXT", "Catatan tambahan"),
        ("created_at", "TIMESTAMP", "Waktu pembuatan record"),
        ("updated_at", "TIMESTAMP", "Waktu update terakhir")
    ]
    add_table_to_doc(doc, ["Field", "Tipe", "Keterangan"], paket_data)
    
    # Tabel Notifikasi
    add_heading(doc, "6.5 Tabel: notifikasi", 2)
    notif_data = [
        ("id", "INT", "Primary Key, Auto Increment"),
        ("penghuni_id", "INT", "Foreign Key ke tabel penghuni"),
        ("paket_id", "INT", "Foreign Key ke tabel paket"),
        ("pesan", "TEXT", "Isi pesan notifikasi"),
        ("is_read", "BOOLEAN", "Status baca/belum baca"),
        ("created_at", "TIMESTAMP", "Waktu notifikasi dibuat")
    ]
    add_table_to_doc(doc, ["Field", "Tipe", "Keterangan"], notif_data)
    
    doc.add_page_break()
    
    # ========== 7. STRUKTUR FOLDER & FILE ==========
    add_heading(doc, "7. STRUKTUR FOLDER & FILE", 1)
    
    add_normal_text(doc,
        "Aplikasi terorganisir dalam struktur folder yang jelas dan mudah dipahami untuk "
        "memudahkan maintenance dan development:")
    
    add_heading(doc, "7.1 Root Level Files", 2)
    root_files = [
        "index.php - Halaman utama/landing page",
        "login.php - Halaman login pengguna",
        "logout.php - Script proses logout",
        "dashboard.php - Dashboard utama setelah login",
        "profile.php - Halaman profil pengguna",
        "forbidden.php - Halaman error 403 (akses terlarang)",
        "database.sql - Database schema dan data awal"
    ]
    for f in root_files:
        add_bullet_point(doc, f)
    
    add_heading(doc, "7.2 Folder Utama", 2)
    
    folders = {
        "config/": [
            "database.php - Konfigurasi koneksi MySQL",
            "session.php - Session management & autentikasi",
            "pagination.php - Konfigurasi pagination"
        ],
        "includes/": [
            "header.php - Header template",
            "navbar.php - Navigation bar",
            "footer.php - Footer template"
        ],
        "modules/": [
            "penghuni/ - Modul manajemen data penghuni",
            "paket/ - Modul manajemen paket",
            "notifikasi/ - Modul sistem notifikasi"
        ],
        "admin/": [
            "index.php - Admin panel utama",
            "users.php - Daftar pengguna",
            "users_add.php - Tambah pengguna",
            "users_edit.php - Edit pengguna",
            "users_delete.php - Hapus pengguna"
        ],
        "assets/": [
            "css/style.css - Stylesheet custom",
            "js/script.js - JavaScript custom",
            "images/ - Folder untuk gambar"
        ],
        "api/": [
            "get-occupied-lokers.php - API endpoint untuk loker terisi"
        ]
    }
    
    for folder, items in folders.items():
        p = doc.add_paragraph(f"📁 {folder}", style='List Bullet')
        for item in items:
            doc.add_paragraph(item, style='List Bullet 2')
    
    add_heading(doc, "7.3 Total File Organization", 2)
    
    org_data = [
        ("PHP Files", "25+"),
        ("Config Files", "3"),
        ("Template Files", "3"),
        ("CSS Files", "1"),
        ("JavaScript Files", "1"),
        ("SQL Files", "1"),
        ("Documentation Files", "10+"),
        ("Total Files", "50+")
    ]
    add_table_to_doc(doc, ["Tipe", "Jumlah"], org_data)
    
    doc.add_page_break()
    
    # ========== 8. PANDUAN INSTALASI ==========
    add_heading(doc, "8. PANDUAN INSTALASI", 1)
    
    add_heading(doc, "8.1 Prasyarat (Requirements)", 2)
    requirements = [
        "XAMPP (Apache + MySQL + PHP)",
        "PHP 7.4 atau lebih tinggi",
        "MySQL 5.7 atau lebih tinggi",
        "Text Editor atau IDE (VS Code, PhpStorm, dll)",
        "Browser modern (Chrome, Firefox, Edge, Safari)"
    ]
    for req in requirements:
        add_bullet_point(doc, req)
    
    add_heading(doc, "8.2 Langkah Instalasi", 2)
    
    steps = [
        ("Ekstrak file", "Ekstrak folder 'sipap' ke dalam folder 'C:\\xampp\\htdocs\\'"),
        ("Start XAMPP", "Buka XAMPP Control Panel dan start Apache dan MySQL"),
        ("Import Database", "Buka phpMyAdmin (http://localhost/phpmyadmin), "
         "buat database baru 'sipap_db', lalu import file 'database.sql'"),
        ("Akses Aplikasi", "Buka browser dan ketik: http://localhost/sipap"),
        ("Login", "Gunakan kredensial default untuk login sebagai admin")
    ]
    
    for i, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title}: ")
        run.bold = True
        p.add_run(desc)
    
    add_heading(doc, "8.3 Kredensial Default", 2)
    
    cred_data = [
        ("Admin", "admin", "admin123"),
        ("Resepsionis", "resepsionis", "resepsionis123"),
        ("Penghuni", "penghuni1", "penghuni123")
    ]
    add_table_to_doc(doc, ["Role", "Username", "Password"], cred_data)
    
    add_normal_text(doc, 
        "⚠️ Penting: Ganti semua password default setelah instalasi berhasil!")
    
    doc.add_page_break()
    
    # ========== 9. PANDUAN PENGGUNAAN ==========
    add_heading(doc, "9. PANDUAN PENGGUNAAN", 1)
    
    add_heading(doc, "9.1 Untuk Admin", 2)
    admin_guide = [
        "Login ke sistem menggunakan akun admin",
        "Akses Admin Panel dari menu dropdown profil",
        "Kelola pengguna: tambah, edit, hapus pengguna",
        "Kelola data penghuni: tambah, edit, hapus penghuni",
        "Monitor semua paket yang masuk dan keluar",
        "Lihat statistik dan ringkasan di dashboard"
    ]
    for guide in admin_guide:
        add_bullet_point(doc, guide)
    
    add_heading(doc, "9.2 Untuk Resepsionis", 2)
    resep_guide = [
        "Login ke sistem menggunakan akun resepsionis",
        "Klik 'Terima Paket Baru' saat kurir datang",
        "Isi form dengan detail paket (pengirim, ekspedisi, deskripsi)",
        "Pilih unit/penghuni penerima dari dropdown",
        "Masukkan nomor loker tempat menyimpan paket",
        "Klik 'Simpan' untuk mencatat paket",
        "Sistem otomatis mengirim notifikasi ke penghuni",
        "Monitor paket yang sedang disimpan di loker",
        "Update status ketika penghuni mengambil paket"
    ]
    for guide in resep_guide:
        add_bullet_point(doc, guide)
    
    add_heading(doc, "9.3 Untuk Penghuni", 2)
    penghuni_guide = [
        "Login ke sistem menggunakan akun penghuni",
        "Lihat notifikasi bell di navbar untuk paket baru",
        "Klik notifikasi untuk melihat detail paket dan nomor loker",
        "Ambil paket dari loker sesuai nomor yang diberikan",
        "Edit profil jika ingin update data pribadi",
        "Lihat riwayat paket yang sudah diambil"
    ]
    for guide in penghuni_guide:
        add_bullet_point(doc, guide)
    
    doc.add_page_break()
    
    # ========== 10. KEAMANAN SISTEM ==========
    add_heading(doc, "10. KEAMANAN SISTEM", 1)
    
    add_heading(doc, "10.1 Implementasi Keamanan", 2)
    
    security_features = {
        "Password Hashing": [
            "Menggunakan bcrypt untuk hashing password",
            "Password tidak pernah disimpan dalam plain text",
            "Salt otomatis pada setiap hash untuk security tambahan"
        ],
        "SQL Injection Prevention": [
            "Menggunakan prepared statements dengan parameterized queries",
            "Input validation di semua form",
            "Sanitasi data sebelum dimasukkan ke database"
        ],
        "XSS (Cross-Site Scripting) Prevention": [
            "Output encoding untuk semua user input",
            "Htmlspecialchars() untuk mencegah script injection",
            "Content Security Policy ready"
        ],
        "Session Management": [
            "Session timeout configuration",
            "Unique session ID untuk setiap user",
            "Session data disimpan di server, bukan di client"
        ],
        "Access Control": [
            "RBAC (Role-Based Access Control)",
            "Direct access protection untuk halaman admin",
            "Permission check di setiap halaman yang restricted",
            "403 Forbidden page untuk akses terlarang"
        ],
        "Data Validation": [
            "Server-side validation di semua form",
            "Input length checking",
            "Data type validation",
            "Required field validation"
        ]
    }
    
    for feature, details in security_features.items():
        add_heading(doc, f"• {feature}", 3)
        for detail in details:
            add_bullet_point(doc, detail)
    
    add_heading(doc, "10.2 Best Practices yang Diikuti", 2)
    
    practices = [
        "Never trust user input - selalu validasi dan sanitasi",
        "Principle of Least Privilege - users hanya dapat akses yang perlu",
        "Error handling yang tidak mengungkap detail sistem",
        "Logging untuk audit trail",
        "Regular backup database",
        "HTTPS ready (saat hosting, gunakan SSL certificate)"
    ]
    for practice in practices:
        add_bullet_point(doc, practice)
    
    doc.add_page_break()
    
    # ========== 11. KESIMPULAN ==========
    add_heading(doc, "11. KESIMPULAN", 1)
    
    add_normal_text(doc,
        "AureliaBox adalah solusi lengkap untuk manajemen paket apartemen premium yang "
        "modern, efisien, dan user-friendly. Dengan implementasi teknologi terkini dan "
        "best practices dalam web development, aplikasi ini telah berhasil mencapai semua "
        "tujuan yang ditetapkan.")
    
    add_heading(doc, "11.1 Pencapaian", 2)
    achievements = [
        "✅ Sistem manajemen paket terintegrasi dan lengkap",
        "✅ 3 role dengan permission control yang tepat",
        "✅ Notifikasi real-time kepada penghuni",
        "✅ Dashboard monitoring untuk semua role",
        "✅ Keamanan tingkat enterprise",
        "✅ User interface yang responsive dan intuitif",
        "✅ Database yang terstruktur dengan baik",
        "✅ Dokumentasi lengkap dan mudah dipahami",
        "✅ Production ready dan siap di-deploy"
    ]
    for achievement in achievements:
        add_bullet_point(doc, achievement)
    
    add_heading(doc, "11.2 Rekomendasi ke Depan", 2)
    
    recommendations = [
        "Implementasi Two-Factor Authentication (2FA) untuk keamanan tambahan",
        "Penambahan fitur reporting dan analytics yang lebih mendalam",
        "Integrasi dengan sistem paging atau SMS notification",
        "Mobile app native untuk penghuni (iOS & Android)",
        "Implementasi real-time notification dengan WebSocket",
        "Integrasi dengan payment gateway untuk layanan tambahan",
        "Automation dengan background jobs untuk tugas periodik",
        "Monitoring dan logging dengan tools seperti ELK Stack"
    ]
    for rec in recommendations:
        add_bullet_point(doc, rec)
    
    add_heading(doc, "11.3 Catatan Akhir", 2)
    add_normal_text(doc,
        "Aplikasi AureliaBox telah dikembangkan dengan standar profesional dan siap untuk "
        "deployment ke environment production. Semua fitur telah diuji dan berjalan dengan "
        "baik. Tim dapat melanjutkan dengan fase hosting dan deployment ke server.")
    
    add_normal_text(doc,
        f"Laporan ini dibuat pada: {datetime.now().strftime('%d %B %Y pukul %H:%M:%S')}")
    
    # Save dokumen
    output_path = r"c:\xampp\htdocs\sipap\LAPORAN_AKHIR_AureliaBox.docx"
    doc.save(output_path)
    print(f"✅ Laporan berhasil dibuat: {output_path}")
    return output_path

if __name__ == "__main__":
    create_report()
